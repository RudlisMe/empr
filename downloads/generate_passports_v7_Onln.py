# -*- coding: utf-8 -*-
"""
Passport generator from Excel/Google Sheets registry and Word template.

V7 Onln update:
- Adds Google Sheets as an optional registry source.
- Accepts a Google Sheets URL as a command-line argument.
- Opens the Google worksheet by gid from the URL.
- Reads Google Sheets headers from row 2.
- Keeps the old local Excel mode when no Google Sheets URL is passed.
- Keeps V6 header aliases and formatting behavior.

Source code is kept ASCII-only where practical. Russian labels and units are
stored with Unicode escape sequences to avoid mojibake in some editors.
"""

from __future__ import annotations

import argparse
import re
import unicodedata
from datetime import datetime
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Sequence
from urllib.parse import parse_qs, urlparse

from docx import Document
from openpyxl import load_workbook

# ===== SETTINGS =====
BASE_DIR = Path(__file__).resolve().parent

TEMPLATE_CANDIDATES = [
    BASE_DIR / "passport_template_v2.docx",
    BASE_DIR / "\u041f\u0430\u0441\u043f\u043e\u0440\u0442_\u0448\u0430\u0431\u043b\u043e\u043d_\u0434\u043b\u044f_Excel_v2.docx",
    BASE_DIR / "\u041f\u0430\u0441\u043f\u043e\u0440\u0442_\u0448\u0430\u0431\u043b\u043e\u043d_\u0434\u043b\u044f_Excel.docx",
]

EXCEL_CANDIDATES = [
    BASE_DIR / "registry.xlsx",
    BASE_DIR / "\u0420\u0435\u0433\u0438\u0441\u0442\u0440\u0430\u0446\u0438\u044f \u0438\u0437\u0434\u0435\u043b\u0438\u0438\u0306 \u0426\u042d\u041e  copy.xlsx",
]

OUTPUT_DIR = BASE_DIR / "Gotovye_pasporta"
SERVICE_ACCOUNT_FILE = BASE_DIR / "google_service_account.json"
GOOGLE_HEADER_ROW = 2

# If True, passports are generated only for rows where the status column is empty.
GENERATE_ONLY_NOT_READY = False

# If True, the script writes the output file name to the status column in Excel.
# Keep False until you check generated documents.
UPDATE_EXCEL_STATUS = False

# Canonical Excel headers, written via Unicode escapes.
H_ORDER_NO = "\u041d\u043e\u043c\u0435\u0440 \u0437\u0430\u043a\u0430\u0437\u0430"
H_PRODUCT_TYPE = "\u0422\u0438\u043f \u0438\u0437\u0434\u0435\u043b\u0438\u044f"
H_FACTORY_NO = "\u0417\u0430\u0432.\u2116"
H_IP = "IP"
H_CURRENT = "\u0422\u043e\u043a, \u0410"
H_VOLTAGE = "\u041d\u0430\u043f\u0440\u044f\u0436\u0435\u043d\u0438\u0435, \u0412"
H_SIZE = "\u0420\u0430\u0437\u043c\u0435\u0440"
H_WEIGHT = "\u0412\u0435\u0441"
H_STATUS = "\u041f\u0430\u0441\u043f\u043e\u0440\u0442 \u0433\u043e\u0442\u043e\u0432"

# Additional accepted headers for changed registry versions.
HEADER_ALIASES: Dict[str, List[str]] = {
    H_ORDER_NO: [
        H_ORDER_NO,
    ],
    H_PRODUCT_TYPE: [
        H_PRODUCT_TYPE,
    ],
    H_FACTORY_NO: [
        H_FACTORY_NO,
        "\u0417\u0430\u0432. \u2116",
        "\u0417\u0430\u0432 \u2116",
        "\u0417\u0430\u0432\u043e\u0434\u0441\u043a\u043e\u0439 \u043d\u043e\u043c\u0435\u0440",
        "\u0417\u0430\u0432.\u043d\u043e\u043c\u0435\u0440",
    ],
    H_IP: [
        H_IP,
    ],
    H_CURRENT: [
        H_CURRENT,
    ],
    H_VOLTAGE: [
        H_VOLTAGE,
    ],
    H_SIZE: [
        H_SIZE,
        "\u0413\u0430\u0431\u0430\u0440\u0438\u0442\u044b",
        "\u0413\u0430\u0431\u0430\u0440\u0438\u0442\u043d\u044b\u0439 \u0440\u0430\u0437\u043c\u0435\u0440",
        "\u0413\u0430\u0431\u0430\u0440\u0438\u0442\u043d\u044b\u0439 \u0440\u0430\u0437\u043c\u0435\u0440, \u0412\u0445\u0428\u0445\u0413, \u043c\u043c",
        "\u0413\u0430\u0431\u0430\u0440\u0438\u0442\u043d\u044b\u0435 \u0440\u0430\u0437\u043c\u0435\u0440\u044b",
    ],
    H_WEIGHT: [
        H_WEIGHT,
        "\u0412\u0435\u0441, \u043a\u0433",
        "\u041c\u0430\u0441\u0441\u0430",
        "\u041c\u0430\u0441\u0441\u0430, \u043a\u0433",
    ],
    H_STATUS: [
        H_STATUS,
    ],
}

# Russian visible text/units, written via Unicode escapes.
RU_NA = "\u041d/\u041f"
RU_AMPERE = "\u0410"
RU_VOLT = "\u0412"
RU_KG = "\u043a\u0433"
RU_WEIGHT_WORD = "\u0412\u0435\u0441"

REQUIRED_HEADERS = [
    H_ORDER_NO, H_PRODUCT_TYPE, H_FACTORY_NO, H_IP,
    H_CURRENT, H_VOLTAGE, H_SIZE, H_WEIGHT,
]


def first_existing(paths: Iterable[Path]) -> Optional[Path]:
    for path in paths:
        if path.exists():
            return path
    return None


def is_google_sheet_url(value: str) -> bool:
    """Return True when a command-line source looks like a Google Sheets URL."""
    return value.startswith("http") and "docs.google.com/spreadsheets" in value


def extract_gid_from_url(sheet_url: str) -> int:
    """Extract worksheet gid from a Google Sheets URL."""
    parsed = urlparse(sheet_url)
    query_gid = parse_qs(parsed.query).get("gid", [None])[0]
    fragment_gid = parse_qs(parsed.fragment).get("gid", [None])[0]
    gid = query_gid or fragment_gid

    if gid is None:
        raise ValueError(
            "Google Sheets URL does not contain gid. "
            "Open the needed worksheet and copy the URL with gid=..."
        )
    try:
        return int(gid)
    except ValueError as exc:
        raise ValueError(f"Invalid Google Sheets gid value: {gid}") from exc


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Generate passports from a local Excel registry or a Google Sheets registry."
    )
    parser.add_argument(
        "source",
        nargs="?",
        help=(
            "Optional source. Pass a Google Sheets URL for online mode, "
            "or an Excel file path for local mode. If omitted, local Excel candidates are used."
        ),
    )
    parser.add_argument(
        "--sheet-url",
        dest="sheet_url",
        help="Google Sheets URL. This has priority over the positional source argument.",
    )
    parser.add_argument(
        "--excel",
        dest="excel_path",
        help="Path to a local Excel registry file. This has priority over local Excel candidates.",
    )
    return parser.parse_args()


def normalize_header(value) -> str:
    """Normalize Excel headers to ignore case, extra spaces, and punctuation spacing."""
    text = str(value or "").replace("\xa0", " ").strip()
    text = unicodedata.normalize("NFKC", text)
    text = text.replace("\u0451", "\u0435").replace("\u0401", "\u0435")
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"\s*([,.;:№/\\-])\s*", r"\1", text)
    return text.lower()


def get_aliases(header: str) -> List[str]:
    """Return accepted Excel header names for a canonical field."""
    aliases = HEADER_ALIASES.get(header, [header])
    result: List[str] = []
    for alias in aliases:
        if alias not in result:
            result.append(alias)
    return result


def alias_norms(header: str) -> List[str]:
    """Return normalized aliases for a canonical field."""
    return [normalize_header(alias) for alias in get_aliases(header)]


def find_header_row(ws, required_headers: Iterable[str], max_scan_rows: int = 30) -> int:
    """Find the Excel header row by required logical columns and their aliases."""
    required = list(required_headers)
    for row_idx in range(1, min(ws.max_row, max_scan_rows) + 1):
        row_values = {normalize_header(cell.value) for cell in ws[row_idx] if cell.value is not None}
        if all(any(alias in row_values for alias in alias_norms(header)) for header in required):
            return row_idx

    missing_info = []
    # Build this from the row with the largest number of matched required fields.
    best_row = 1
    best_score = -1
    best_values = set()
    for row_idx in range(1, min(ws.max_row, max_scan_rows) + 1):
        row_values = {normalize_header(cell.value) for cell in ws[row_idx] if cell.value is not None}
        score = sum(any(alias in row_values for alias in alias_norms(header)) for header in required)
        if score > best_score:
            best_row = row_idx
            best_score = score
            best_values = row_values
    for header in required:
        if not any(alias in best_values for alias in alias_norms(header)):
            missing_info.append(" / ".join(get_aliases(header)))
    details = "; ".join(missing_info)
    raise RuntimeError(
        f"Excel header row was not found. Best candidate row: {best_row}. "
        f"Missing logical columns: {details}"
    )


def build_header_map(ws, header_row: int) -> Dict[str, int]:
    """Return a map: normalized actual Excel header -> column number."""
    result = {}
    for cell in ws[header_row]:
        if cell.value is not None:
            result[normalize_header(cell.value)] = cell.column
    return result


def build_header_map_from_values(header_values: Sequence[object]) -> Dict[str, int]:
    """Return a map: normalized header -> 1-based column number."""
    result = {}
    for index, value in enumerate(header_values, start=1):
        if value is not None and str(value).strip() != "":
            result[normalize_header(value)] = index
    return result


def check_required_columns(header_map: Dict[str, int], required_headers: Iterable[str]) -> None:
    """Validate that all required logical columns are present in a header map."""
    missing = []
    for header in required_headers:
        if resolve_column(header_map, header) is None:
            missing.append(" / ".join(get_aliases(header)))
    if missing:
        raise RuntimeError("Missing required columns: " + "; ".join(missing))


def resolve_column(header_map: Dict[str, int], header: str) -> Optional[int]:
    """Resolve a canonical field to an Excel column using aliases."""
    for alias in alias_norms(header):
        col_idx = header_map.get(alias)
        if col_idx:
            return col_idx
    return None


def is_empty(value) -> bool:
    return value is None or str(value).strip() == ""


def is_zero(value) -> bool:
    if is_empty(value):
        return False
    try:
        return float(str(value).replace(",", ".").strip()) == 0
    except ValueError:
        return False


def format_value(value) -> str:
    """Format Excel values without units."""
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    if isinstance(value, datetime):
        return value.strftime("%d.%m.%Y")
    return str(value).strip()


def format_raw_na(value) -> str:
    """Format current/voltage raw value. Zero becomes N/P, unit is not added."""
    if is_empty(value):
        return ""
    if is_zero(value):
        return RU_NA
    return format_value(value)


def format_display(value, unit: str, zero_as_na: bool = False, empty_as: str = "") -> str:
    """Format value for display placeholder. Unit is added only when applicable."""
    if is_empty(value):
        return empty_as
    if zero_as_na and is_zero(value):
        return RU_NA
    return f"{format_value(value)}{unit}"


def sanitize_filename(text: str) -> str:
    """Remove characters that are invalid in Windows file names."""
    text = re.sub(r'[\\/:*?"<>|]+', "_", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text[:180] or "passport"


def cleanup_generated_text(text: str) -> str:
    """Clean common duplicate artifacts from older template/script combinations."""
    # Example: duplicate Russian word/unit for weight -> single word/unit.
    text = text.replace(f"{RU_WEIGHT_WORD} {RU_WEIGHT_WORD} ", f"{RU_WEIGHT_WORD} ")
    text = text.replace(f"{RU_KG}{RU_KG}", RU_KG)
    # Keep cleanup conservative: avoid changing voltage/current duplicate units globally.
    return text


def replace_in_paragraph(paragraph, replacements: Dict[str, str]) -> None:
    """
    Replace placeholders inside a paragraph.
    The first pass keeps formatting when a placeholder is inside one run.
    The fallback pass handles placeholders split across multiple runs.
    """
    if not paragraph.runs:
        return

    for run in paragraph.runs:
        original = run.text
        new_text = original
        for key, value in replacements.items():
            new_text = new_text.replace(key, value)
        new_text = cleanup_generated_text(new_text)
        if new_text != original:
            run.text = new_text

    combined = "".join(run.text for run in paragraph.runs)
    if not any(key in combined for key in replacements):
        cleaned = cleanup_generated_text(combined)
        if cleaned != combined:
            paragraph.runs[0].text = cleaned
            for run in paragraph.runs[1:]:
                run.text = ""
        return

    new_combined = combined
    for key, value in replacements.items():
        new_combined = new_combined.replace(key, value)
    new_combined = cleanup_generated_text(new_combined)

    first_run = paragraph.runs[0]
    for run in paragraph.runs[1:]:
        run.text = ""
    first_run.text = new_combined


def iter_all_paragraphs(doc: Document):
    """Yield paragraphs from body, tables, headers, and footers."""
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
                for nested_table in cell.tables:
                    for nested_row in nested_table.rows:
                        for nested_cell in nested_row.cells:
                            for p in nested_cell.paragraphs:
                                yield p
    for section in doc.sections:
        for container in (section.header, section.footer):
            for p in container.paragraphs:
                yield p
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            yield p


def fill_docx_template(template_path: Path, output_path: Path, replacements: Dict[str, str]) -> None:
    doc = Document(template_path)
    for paragraph in iter_all_paragraphs(doc):
        replace_in_paragraph(paragraph, replacements)
    doc.save(output_path)


def ask_start_row(default_start_row: int, max_row: int, source_label: str = "registry") -> int:
    """
    Ask user for the registry row where passport generation should start.
    Empty input means default_start_row.
    """
    prompt = (
        f"Start {source_label} row for generation "
        f"[press Enter for {default_start_row}]: "
    )
    while True:
        user_input = input(prompt).strip()
        if user_input == "":
            return default_start_row
        try:
            start_row = int(user_input)
        except ValueError:
            print("Please enter a valid row number, for example 20.")
            continue
        if start_row < default_start_row:
            print(f"Row must be {default_start_row} or greater.")
            continue
        if start_row > max_row:
            print(f"Row is greater than the last Excel row ({max_row}).")
            continue
        return start_row


def load_google_sheet(sheet_url: str):
    """Open a Google Sheets worksheet by gid and return worksheet plus all values."""
    if not SERVICE_ACCOUNT_FILE.exists():
        raise FileNotFoundError(
            f"Google service account file was not found: {SERVICE_ACCOUNT_FILE}"
        )

    try:
        import gspread
    except ImportError as exc:
        raise ImportError(
            "Google Sheets mode requires packages: gspread google-auth. "
            "Install them with: pip install gspread google-auth"
        ) from exc

    gid = extract_gid_from_url(sheet_url)
    gc = gspread.service_account(filename=str(SERVICE_ACCOUNT_FILE))
    spreadsheet = gc.open_by_url(sheet_url)
    worksheet = spreadsheet.get_worksheet_by_id(gid)
    if worksheet is None:
        raise RuntimeError(f"Worksheet with gid={gid} was not found in this spreadsheet.")
    values = worksheet.get_all_values()
    return spreadsheet, worksheet, gid, values


def main() -> None:
    args = parse_args()

    template_path = first_existing(TEMPLATE_CANDIDATES)
    if template_path is None:
        raise FileNotFoundError(
            "Word template was not found. Use passport_template_v2.docx near this script."
        )

    source_arg = args.sheet_url or args.excel_path or args.source
    use_google = bool(args.sheet_url) or (bool(source_arg) and is_google_sheet_url(str(source_arg)))

    OUTPUT_DIR.mkdir(exist_ok=True)

    wb = None
    ws = None
    worksheet = None
    google_values: List[List[str]] = []
    source_label = "Excel"

    if use_google:
        sheet_url = str(source_arg)
        spreadsheet, worksheet, gid, google_values = load_google_sheet(sheet_url)
        header_row = GOOGLE_HEADER_ROW
        if len(google_values) < header_row:
            raise RuntimeError(f"Google worksheet has no header row {header_row}.")
        header_map = build_header_map_from_values(google_values[header_row - 1])
        check_required_columns(header_map, REQUIRED_HEADERS)
        max_row = max(len(google_values), header_row)
        source_label = "Google Sheets"

        print(f"Google spreadsheet: {spreadsheet.title}")
        print(f"Google worksheet: {worksheet.title} (gid={gid})")
        print(f"Google Sheets header row used: {header_row}")

        def get_cell_value(row_idx: int, header: str):
            col_idx = resolve_column(header_map, header)
            if not col_idx or row_idx < 1 or row_idx > len(google_values):
                return None
            row_values = google_values[row_idx - 1]
            if col_idx > len(row_values):
                return None
            return row_values[col_idx - 1]

        def get_status_value(row_idx: int, status_col: Optional[int]):
            if not status_col or row_idx < 1 or row_idx > len(google_values):
                return None
            row_values = google_values[row_idx - 1]
            if status_col > len(row_values):
                return None
            return row_values[status_col - 1]

    else:
        if args.excel_path:
            excel_path = Path(args.excel_path).expanduser()
        elif source_arg:
            excel_path = Path(str(source_arg)).expanduser()
        else:
            excel_path = first_existing(EXCEL_CANDIDATES)

        if excel_path is None or not excel_path.exists():
            raise FileNotFoundError(
                "Excel registry was not found. Use registry.xlsx near this script, "
                "or pass an Excel file path / Google Sheets URL when starting the script."
            )

        wb = load_workbook(excel_path)
        ws = wb.active
        header_row = find_header_row(ws, REQUIRED_HEADERS)
        header_map = build_header_map(ws, header_row)
        max_row = ws.max_row
        source_label = "Excel"

        print(f"Excel registry: {excel_path}")
        print(f"Excel header row detected: {header_row}")

        def get_cell_value(row_idx: int, header: str):
            col_idx = resolve_column(header_map, header)
            if not col_idx:
                return None
            return ws.cell(row=row_idx, column=col_idx).value

        def get_status_value(row_idx: int, status_col: Optional[int]):
            if not status_col:
                return None
            return ws.cell(row=row_idx, column=status_col).value

    status_col = resolve_column(header_map, H_STATUS)
    default_start_row = header_row + 1
    start_row = ask_start_row(default_start_row, max_row, source_label)

    created = []
    skipped = 0

    size_col = resolve_column(header_map, H_SIZE)
    weight_col = resolve_column(header_map, H_WEIGHT)
    print(f"Size column detected: {size_col}")
    print(f"Weight column detected: {weight_col}")
    print(f"Generating passports from {source_label} row {start_row} to {max_row}...")

    for row_idx in range(start_row, max_row + 1):
        order_no_raw = get_cell_value(row_idx, H_ORDER_NO)
        factory_no_raw = get_cell_value(row_idx, H_FACTORY_NO)
        product_type_raw = get_cell_value(row_idx, H_PRODUCT_TYPE)

        order_no = format_value(order_no_raw)
        factory_no = format_value(factory_no_raw)
        product_type = format_value(product_type_raw)

        if not any([order_no, factory_no, product_type]):
            skipped += 1
            continue

        if GENERATE_ONLY_NOT_READY and status_col:
            status = format_value(get_status_value(row_idx, status_col))
            if status:
                skipped += 1
                continue

        current_raw = get_cell_value(row_idx, H_CURRENT)
        voltage_raw = get_cell_value(row_idx, H_VOLTAGE)
        weight_raw = get_cell_value(row_idx, H_WEIGHT)

        replacements = {
            "{{ORDER_NO}}": order_no,
            "{{PRODUCT_TYPE}}": product_type,
            "{{FACTORY_NO}}": factory_no,
            "{{IP}}": format_value(get_cell_value(row_idx, H_IP)),
            "{{SIZE}}": format_value(get_cell_value(row_idx, H_SIZE)),

            # Raw placeholders: no labels and no units.
            "{{CURRENT}}": format_raw_na(current_raw),
            "{{VOLTAGE}}": format_raw_na(voltage_raw),
            "{{WEIGHT}}": format_value(weight_raw) if not is_empty(weight_raw) else "___",

            # Display placeholders: include units when value is applicable.
            "{{CURRENT_DISPLAY}}": format_display(current_raw, RU_AMPERE, zero_as_na=True),
            "{{VOLTAGE_DISPLAY}}": format_display(voltage_raw, RU_VOLT, zero_as_na=True),
            "{{WEIGHT_DISPLAY}}": format_display(weight_raw, RU_KG, empty_as="___"),
        }

        filename = sanitize_filename(f"Passport_order_{order_no}_factory_{factory_no}.docx")
        output_path = OUTPUT_DIR / filename
        fill_docx_template(template_path, output_path, replacements)
        created.append(output_path.name)

        if UPDATE_EXCEL_STATUS and status_col:
            status_text = f"Created: {output_path.name}"
            if use_google and worksheet is not None:
                worksheet.update_cell(row_idx, status_col, status_text)
            elif ws is not None:
                ws.cell(row=row_idx, column=status_col).value = status_text

    if UPDATE_EXCEL_STATUS and status_col and not use_google and wb is not None:
        wb.save(excel_path)

    print(f"Done. Created passports: {len(created)}")
    print(f"Start row used: {start_row}")
    if skipped:
        print(f"Skipped rows: {skipped}")
    print(f"Output folder: {OUTPUT_DIR}")
    for name in created[:20]:
        print(" -", name)
    if len(created) > 20:
        print(f"... and {len(created) - 20} more")


if __name__ == "__main__":
    main()
